# -*- coding: utf-8 -*-
"""Extrai de uma MIT044 ja pronta o template limpo + os prototipos de formatacao
usados pela skill mit044-especificacao. Rode uma vez para adotar o layout da sua
empresa/cliente.

Uso:
    python extrai_template.py <mit044-referencia.docx> <pasta-assets>
                              [--proto-body PREFIXO] [--proto-label PREFIXO] [--proto-bullet PREFIXO]

Gera:
    <pasta-assets>/template-mit044.docx   esqueleto sem conteudo (capa em branco)
    <pasta-assets>/prototipos.xml         fragmentos de formatacao (3 paragrafos + 3 tabelas)

Os tres paragrafos-prototipo sao detectados automaticamente no documento de origem
(o primeiro paragrafo narrativo, o primeiro rotulo em negrito terminado em ":" e o
primeiro bullet). Use os parametros --proto-* para apontar outro paragrafo por
prefixo de texto quando a deteccao automatica escolher mal.
"""
import sys, os, copy
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.table import Table
from lxml import etree

W_P, W_TBL, W_R, W_T = qn('w:p'), qn('w:tbl'), qn('w:r'), qn('w:t')
SECOES = ['Processo Atual', 'Processo Proposto', 'Parametrizações', 'Execução', 'Customizações']

# cabecalho da primeira linha identifica cada tabela de conteudo (mais robusto que indice)
PROTO_TABELAS = {
    'periodicidade': 'Forma de execução',
    'rotina':        'Rotina (código provisório)',
    'anexos':        'Descrição',
}
BULLET = '•'


def acha_paragrafo(doc, prefixo):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefixo):
            return p._p
    raise SystemExit('protótipo de parágrafo não encontrado: ' + prefixo)


def detecta_prototipos(doc):
    """Descobre no documento os três parágrafos que servirão de molde de formatação.

    body   = primeiro parágrafo narrativo longo (não bullet, não negrito)
    label  = primeiro rótulo em negrito terminado em ':'
    bullet = primeiro parágrafo iniciado pelo caractere de bullet
    A varredura começa depois do heading 'Processo Atual' — antes dele só existe capa.
    """
    achados = {}
    dentro = False
    for p in doc.paragraphs:
        txt = p.text.strip()
        if p.style.name in ('Heading 1', 'Heading 2') and txt == 'Processo Atual':
            dentro = True
            continue
        if not dentro or not txt or p.style.name.startswith('Heading'):
            continue
        negrito = bool(p.runs and p.runs[0].bold)
        if txt.startswith(BULLET):
            achados.setdefault('bullet', p._p)
        elif negrito and txt.endswith(':'):
            achados.setdefault('label', p._p)
        elif not negrito and len(txt) > 80:
            achados.setdefault('body', p._p)
        if len(achados) == 3:
            break
    faltando = [k for k in ('body', 'label', 'bullet') if k not in achados]
    if faltando:
        raise SystemExit(
            'não foi possível detectar os protótipos: %s. Informe-os por prefixo com '
            '--proto-body / --proto-label / --proto-bullet.' % ', '.join(faltando))
    return achados


def acha_tabela(doc, cabecalho):
    for t in doc.tables:
        if t.rows and t.rows[0].cells[0].text.strip() == cabecalho:
            return t._tbl
    raise SystemExit('tabela protótipo não encontrada: ' + cabecalho)


def limpa_valor_celula(cell):
    """Esvazia o VALOR da celula da capa, preservando o rotulo (1o run) e a formatacao."""
    for p in cell.paragraphs:
        runs = p.runs
        if len(runs) >= 2:
            for r in runs[1:]:
                r.text = ''
        elif len(runs) == 1 and ':' in runs[0].text:
            runs[0].text = runs[0].text.split(':', 1)[0] + ': '


def main():
    import argparse
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument('origem', help='MIT044 aprovada usada como doadora de formatação')
    ap.add_argument('destino', help='pasta onde gravar os assets')
    ap.add_argument('--proto-body', help='prefixo do parágrafo narrativo protótipo')
    ap.add_argument('--proto-label', help='prefixo do rótulo em negrito protótipo')
    ap.add_argument('--proto-bullet', help='prefixo do bullet protótipo')
    args = ap.parse_args()
    origem, destino = args.origem, args.destino
    os.makedirs(destino, exist_ok=True)

    doc = docx.Document(origem)
    body = doc.element.body

    # ---------- 1. extrair prototipos ANTES de limpar ----------
    detectados = detecta_prototipos(doc)
    manuais = {'body': args.proto_body, 'label': args.proto_label, 'bullet': args.proto_bullet}
    fragmentos = []
    for nome in ('body', 'label', 'bullet'):
        el = acha_paragrafo(doc, manuais[nome]) if manuais[nome] else detectados[nome]
        fragmentos.append((nome, copy.deepcopy(el)))
        origem_proto = 'informado' if manuais[nome] else 'detectado'
        print('  protótipo %-7s (%s): %r' % (
            nome, origem_proto, ''.join(el.itertext())[:58]))
    for nome, cab in PROTO_TABELAS.items():
        el = copy.deepcopy(acha_tabela(doc, cab))
        fragmentos.append((nome, el))

    partes = ['<protos>']
    for nome, el in fragmentos:
        partes.append('<proto id="%s">' % nome)
        partes.append(etree.tostring(el, encoding='unicode'))
        partes.append('</proto>')
    partes.append('</protos>')
    xml_protos = '\n'.join(partes)
    # valida o XML antes de gravar
    etree.fromstring(xml_protos.encode('utf-8'))
    caminho_protos = os.path.join(destino, 'prototipos.xml')
    with open(caminho_protos, 'w', encoding='utf-8') as fh:
        fh.write(xml_protos)
    print('prototipos.xml:', len(xml_protos), 'bytes |', ', '.join(n for n, _ in fragmentos))

    # ---------- 2. limpar o miolo (paragrafos E tabelas) ----------
    heads = {}
    for el in body.iterchildren():
        if el.tag == W_P:
            p = Paragraph(el, doc)
            txt = p.text.strip()
            if p.style.name in ('Heading 1', 'Heading 2') and txt in SECOES + ['Aceite']:
                heads[txt] = el
    faltando = [s for s in SECOES + ['Aceite'] if s not in heads]
    if faltando:
        raise SystemExit('headings não encontrados no documento de origem: ' + ', '.join(faltando))

    manter = {id(heads[s]) for s in SECOES}
    dentro, rem_p, rem_t = False, 0, 0
    for el in list(body.iterchildren()):
        if el is heads['Processo Atual']:
            dentro = True
            continue
        if el is heads['Aceite']:
            break
        if dentro and id(el) not in manter:
            if el.tag == W_P:
                rem_p += 1
            elif el.tag == W_TBL:
                rem_t += 1
            body.remove(el)
    print('miolo removido: %d parágrafos, %d tabelas' % (rem_p, rem_t))

    # ---------- 3. esvaziar os valores da capa ----------
    capa = doc.tables[0]
    for row in capa.rows:
        for tc in row._tr.tc_lst:
            from docx.table import _Cell
            limpa_valor_celula(_Cell(tc, capa))
    # "Responsavel na TOTVS" da tabela 2 tambem e' dado de pessoa
    dados = doc.tables[1]
    for row in dados.rows:
        txt = row.cells[0].text
        if txt.startswith('Responsável na TOTVS'):
            for p in row.cells[0].paragraphs:
                runs = p.runs
                if len(runs) >= 2:
                    for r in runs[1:]:
                        r.text = ''
                elif len(runs) == 1:
                    runs[0].text = 'Responsável na TOTVS: '

    caminho_tpl = os.path.join(destino, 'template-mit044.docx')
    doc.save(caminho_tpl)
    print('template-mit044.docx:', os.path.getsize(caminho_tpl), 'bytes')

    # ---------- 4. conferencia rapida ----------
    novo = docx.Document(caminho_tpl)
    ordem = [p.style.name + ': ' + p.text.strip()
             for p in novo.paragraphs
             if p.text.strip() and p.style.name in ('Title', 'Heading 1', 'Heading 2')]
    print('esqueleto:', ' | '.join(ordem))
    print('tabelas restantes:', len(novo.tables))
    toc = [e.text for e in novo.element.body.iter()
           if e.tag == qn('w:instrText') and e.text and 'TOC' in e.text]
    print('campo TOC preservado:', bool(toc))
    print('capa (linha 1):', repr(novo.tables[0].rows[1].cells[0].text))


if __name__ == '__main__':
    main()
