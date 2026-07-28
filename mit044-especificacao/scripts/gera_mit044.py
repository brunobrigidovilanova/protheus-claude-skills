# -*- coding: utf-8 -*-
"""Gera um documento MIT044 - Especificacao da Customizacao (padrao TOTVS) em .docx
a partir de um JSON de conteudo, usando o template e os prototipos de formatacao
da skill mit044-especificacao.

Uso:
    python gera_mit044.py <conteudo.json> [--saida PASTA] [--template DOCX] [--force]

O nome do arquivo de saida e montado como:
    [<tag_cliente>] - Especificacao da Customizacao - MIT044 - <titulo>.docx
"""
import sys, os, re, json, copy, shutil, argparse
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from lxml import etree

AQUI = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.normpath(os.path.join(AQUI, '..', 'assets'))
TEMPLATE_PADRAO = os.path.join(ASSETS, 'template-mit044.docx')
PROTOTIPOS = os.path.join(ASSETS, 'prototipos.xml')

W_P, W_TBL, W_R, W_T = qn('w:p'), qn('w:tbl'), qn('w:r'), qn('w:t')
BULLET = '•  '          # bullet literal + 2 espacos (padrao dos documentos)
VAZIO, MARCADO = '☐', '☒'

# o template oficial traz placeholders {{campo}} na capa e paragrafos de orientacao
# entre <>; os dois somem do documento gerado
PLACEHOLDER = re.compile(r'\{\{[^{}]*\}\}')
INSTRUCAO = re.compile(r'^<.*>\.?$', re.S)

SECOES = ['Processo Atual', 'Processo Proposto', 'Parametrizações', 'Execução', 'Customizações']

# rotulo na tabela de capa -> chave no JSON
CAPA = [
    ('Nome do cliente',            'nome_cliente'),
    ('Código de cliente',          'codigo_cliente'),
    ('Nome do projeto',            'nome_projeto'),
    ('Código do projeto',          'codigo_projeto'),
    ('Segmento cliente',           'segmento_cliente'),
    ('Unidade TOTVS',              'unidade_totvs'),
    ('Data',                       'data'),
    ('Proposta comercial',         'proposta_comercial'),
    ('Gerente/Coordenador TOTVS',  'gerente_totvs'),
    ('Gerente/Coordenador cliente', 'gerente_cliente'),
]
DADOS_CUST = [
    ('Qtd. Horas',            'qtd_horas'),
    ('Responsável no Cliente', 'responsavel_cliente'),
    ('Responsável na TOTVS',  'responsavel_totvs'),
]
CRITICIDADE = {'alto': 'Alto Impacto', 'medio': 'Médio Impacto', 'baixo': 'Baixo Impacto'}

# colunas da tabela "Historico de Versoes" (secao inicial do template oficial)
HISTORICO = [('data', 'data'), ('versao', 'versão'), ('autor', 'autor'),
             ('descricao', 'descrição')]

# labels fixos dos blocos, na ordem em que aparecem no documento
BLOCOS_EXECUCAO = [
    ('objetivos',       'Objetivos do negócio:',    'bullet'),
    ('fluxo',           'Fluxo do processo:',       'num'),
    ('premissas',       'Premissas e Restrições:',  'bullet'),
    ('plano_teste',     'Plano de teste e cenários esperados:', 'bullet'),
    ('rastreabilidade', 'Rastreabilidade / dependência com outra MIT044:', 'p'),
]
LABELS_CUSTOMIZACOES = [
    'Periodicidade de execução:', 'Onde será executada:', 'Funcionalidades:',
    'Premissas e restrições técnicas:', 'Protótipo de tela:', 'Anexos:',
]
TODOS_LABELS = {l for _, l, _ in BLOCOS_EXECUCAO} | set(LABELS_CUSTOMIZACOES)


class ErroConteudo(Exception):
    pass


# ------------------------------------------------------------------ prototipos
def carrega_prototipos():
    if not os.path.isfile(PROTOTIPOS):
        raise ErroConteudo('prototipos.xml não encontrado em ' + ASSETS)
    root = etree.parse(PROTOTIPOS).getroot()
    return {p.get('id'): p[0] for p in root}


def clona_paragrafo(proto, texto):
    """Clona um <w:p> protótipo trocando o texto e preservando pPr/rPr."""
    p = copy.deepcopy(proto)
    for tag in ('w:hyperlink', 'w:bookmarkStart', 'w:bookmarkEnd', 'w:proofErr'):
        for el in p.findall(qn(tag)):
            p.remove(el)
    runs = p.findall(W_R)
    if not runs:
        raise ErroConteudo('protótipo de parágrafo sem run')
    primeiro = runs[0]
    for r in runs[1:]:
        p.remove(r)
    ts = primeiro.findall(W_T)
    for t in ts[1:]:
        primeiro.remove(t)
    if not ts:
        t = etree.SubElement(primeiro, W_T)
        ts = [t]
    ts[0].text = texto
    ts[0].set(qn('xml:space'), 'preserve')
    return p


def texto_celula(tc, texto, doc):
    """Escreve texto em uma celula de tabela preservando a formatacao do 1o run."""
    cell = _Cell(tc, doc)
    ps = tc.findall(W_P)
    for extra in ps[1:]:
        tc.remove(extra)
    p = ps[0]
    runs = p.findall(W_R)
    if not runs:
        r = etree.SubElement(p, W_R)
        runs = [r]
    primeiro = runs[0]
    for r in runs[1:]:
        p.remove(r)
    ts = primeiro.findall(W_T)
    for t in ts[1:]:
        primeiro.remove(t)
    if not ts:
        t = etree.SubElement(primeiro, W_T)
        ts = [t]
    ts[0].text = texto
    ts[0].set(qn('xml:space'), 'preserve')


def clona_tabela(proto, linhas, doc):
    """Clona a tabela protótipo ajustando o número de linhas e o conteúdo."""
    tbl = copy.deepcopy(proto)
    trs = tbl.findall(qn('w:tr'))
    if len(linhas) > len(trs):
        modelo = copy.deepcopy(trs[-1])
        for _ in range(len(linhas) - len(trs)):
            tbl.append(copy.deepcopy(modelo))
        trs = tbl.findall(qn('w:tr'))
    for tr in trs[len(linhas):]:
        tbl.remove(tr)
    trs = tbl.findall(qn('w:tr'))
    for tr, linha in zip(trs, linhas):
        tcs = tr.findall(qn('w:tc'))
        for tc, txt in zip(tcs, linha):
            texto_celula(tc, txt, doc)
    return tbl


# ------------------------------------------------------------------ documento
def acha_headings(doc):
    heads = {}
    for el in doc.element.body.iterchildren():
        if el.tag == W_P:
            p = Paragraph(el, doc)
            txt = p.text.strip()
            if p.style.name in ('Heading 1', 'Heading 2') and txt in SECOES + ['Aceite']:
                heads[txt] = el
    faltando = [s for s in SECOES + ['Aceite'] if s not in heads]
    if faltando:
        raise ErroConteudo('template sem os headings: ' + ', '.join(faltando))
    return heads


def limpa_quebras(heads):
    """Tira dos títulos as quebras de linha de uma geração anterior.

    Quando o template é extraído de uma MIT044 já gerada, os `<w:br>` do respiro vêm
    junto nos títulos preservados — sem esta limpeza cada nova geração acrescentaria
    mais uma quebra.
    """
    n = 0
    for el in heads.values():
        for r in list(el.findall(W_R)):
            brs = r.findall(qn('w:br'))
            if not brs:
                continue
            if ''.join(t.text or '' for t in r.findall(W_T)).strip():
                for br in brs:          # run com texto: sai só a quebra
                    r.remove(br)
            else:
                el.remove(r)            # run que só existia para a quebra
            n += len(brs)
    return n


def limpa_miolo(doc, heads):
    """Remove parágrafos E tabelas entre 'Processo Atual' e 'Aceite' (idempotência)."""
    body = doc.element.body
    manter = {id(heads[s]) for s in SECOES}
    dentro, removidos = False, 0
    for el in list(body.iterchildren()):
        if el is heads['Processo Atual']:
            dentro = True
            continue
        if el is heads['Aceite']:
            break
        if dentro and id(el) not in manter:
            body.remove(el)
            removidos += 1
    return removidos


def corrige_sumario(doc):
    """Faz o campo TOC do template montar o sumário em qualquer idioma do Word.

    O template traz `TOC \\t "Heading 1,1,Heading 2,2,..."` — seleção por NOME de
    estilo em inglês. No Word em português os estilos se chamam "Título 1"/"Título 2",
    nada casa e o F9 devolve "Nenhuma entrada de sumário foi encontrada". Trocando
    para `\\o "1-2"` a seleção passa a ser pelo nível de estrutura de tópicos; nos
    templates exportados do Google Docs esse nível ainda precisa ser declarado nos
    estilos de título, que vêm sem ele.
    """
    ajustes = 0
    for style in doc.styles.element.findall(qn('w:style')):
        nome = style.find(qn('w:name'))
        if nome is None:
            continue
        m = re.match(r'heading ([1-9])$', (nome.get(qn('w:val')) or '').lower())
        if not m:
            continue
        pPr = style.get_or_add_pPr()
        if pPr.find(qn('w:outlineLvl')) is None:
            lvl = etree.SubElement(pPr, qn('w:outlineLvl'))
            lvl.set(qn('w:val'), str(int(m.group(1)) - 1))
            ajustes += 1
    for instr in doc.element.body.iter(qn('w:instrText')):
        if instr.text and 'TOC' in instr.text and '\\t' in instr.text:
            instr.text = ' TOC \\o "1-2" \\h \\u \\z '
            ajustes += 1
    return ajustes


def uniformiza_numeracao(doc, heads):
    """Poe os 5 Heading 2 na MESMA lista numerada.

    Nos documentos de origem, 'Processo Atual' e 'Processo Proposto' costumam usar
    numId=1/ilvl=1 e os tres seguintes numId=5/ilvl=0 — a numeracao sai quebrada
    (a., b., 01., 02., 03.). Só é aplicado com --numeracao-uniforme.
    """
    ref = heads['Parametrizações'].find(qn('w:pPr')).find(qn('w:numPr'))
    if ref is None:
        return 0
    trocados = 0
    for nome in SECOES:
        pPr = heads[nome].find(qn('w:pPr'))
        atual = pPr.find(qn('w:numPr'))
        if atual is not None:
            pPr.remove(atual)
        novo = copy.deepcopy(ref)
        # no OOXML o numPr vem logo depois do pStyle dentro do pPr
        pStyle = pPr.find(qn('w:pStyle'))
        if pStyle is not None:
            pStyle.addnext(novo)
        else:
            pPr.insert(0, novo)
        trocados += 1
    return trocados


def acha_tabela(doc, texto, indice_padrao=None):
    """Localiza uma tabela pelo texto de qualquer célula da 1a linha.

    Localizar por conteúdo (e não por índice) mantém o gerador funcionando quando
    o template ganha ou perde tabelas — foi o que aconteceu quando o "Histórico de
    Versões" entrou entre a capa e os "Dados da Customização".
    """
    for t in doc.tables:
        if not t.rows:
            continue
        for cell in t.rows[0].cells:
            if cell.text.strip().startswith(texto):
                return t
        # a capa tem a 1a linha mesclada e vazia: olha também a linha seguinte
        if len(t.rows) > 1 and t.rows[1].cells[0].text.strip().startswith(texto):
            return t
    if indice_padrao is not None and len(doc.tables) > indice_padrao:
        return doc.tables[indice_padrao]
    return None


def preenche_capa(doc, capa):
    tab = acha_tabela(doc, 'Nome do cliente', 0)
    if tab is None:
        raise ErroConteudo('tabela de capa não encontrada no template')
    achados = set()
    for row in tab.rows:
        for tc in row._tr.tc_lst:
            cell = _Cell(tc, tab)
            txt = cell.text.strip()
            for rotulo, chave in CAPA:
                if txt.startswith(rotulo + ':') and chave not in achados:
                    escreve_valor(cell, rotulo, capa.get(chave, ''))
                    achados.add(chave)
                    break
    faltou = [c for _, c in CAPA if c not in achados]
    if faltou:
        print('  aviso: rótulos não localizados na capa: ' + ', '.join(faltou))

    # tabela "Dados da Customizacao"
    dados = acha_tabela(doc, 'Dados da Customização', 1)
    if dados is None:
        raise ErroConteudo('tabela "Dados da Customização" não encontrada no template')
    for row in dados.rows:
        cell = row.cells[0]
        txt = cell.text.strip()
        for rotulo, chave in DADOS_CUST:
            if txt.startswith(rotulo):
                escreve_valor(cell, rotulo, capa.get(chave, ''))
                break
        if txt.startswith('Extra Projeto'):
            marca_checkbox(cell, {'sim': 'Sim', 'nao': 'Não'}.get(
                str(capa.get('extra_projeto', '')).lower(), ''))
        if txt.startswith('Criticidade'):
            marca_checkbox(cell, CRITICIDADE.get(
                str(capa.get('criticidade', '')).lower(), ''))


def escreve_valor(cell, rotulo, valor):
    """Reescreve a célula como "rótulo: valor", preservando o run do rótulo (negrito).

    No template oficial o rótulo às vezes vem grudado ao placeholder ({{cliente}}) no
    mesmo run — por isso o rótulo é reescrito por inteiro em vez de ter só o valor
    trocado. O valor vai num run próprio, sem negrito.
    """
    valor = '' if valor is None else str(valor).strip()
    for p in cell.paragraphs:
        runs = p.runs
        if not runs:
            continue
        runs[0].text = rotulo + ': '
        for r in runs[1:]:
            r.text = ''
        if valor:
            alvo = runs[1] if len(runs) >= 2 else clona_run(p, runs[0])
            alvo.text = valor
            alvo.bold = False
        return


def clona_run(p, modelo):
    """Acrescenta ao parágrafo um run com a mesma formatação do modelo."""
    novo = copy.deepcopy(modelo._r)
    for t in novo.findall(W_T):
        novo.remove(t)
    modelo._r.addnext(novo)
    from docx.text.run import Run
    return Run(novo, p)


def limpa_placeholders(doc):
    """Apaga os {{campo}} do template que nenhum valor do JSON sobrescreveu."""
    n = 0
    for t in doc.element.body.iter(W_T):
        if t.text and PLACEHOLDER.search(t.text):
            t.text = PLACEHOLDER.sub('', t.text).strip()
            n += 1
    return n


def remove_instrucoes(doc):
    """Remove os parágrafos de orientação do template (texto entre < e >).

    São as instruções do documento oficial ("<Neste local deve ser descrito...>").
    As que ficam entre "Processo Atual" e "Aceite" já saem em limpa_miolo; esta
    varredura pega as de fora, como a do Histórico de Versões.
    """
    body = doc.element.body
    n = 0
    for el in list(body.iterchildren()):
        if el.tag != W_P:
            continue
        txt = ''.join(el.itertext()).strip()
        if txt and INSTRUCAO.match(txt):
            body.remove(el)
            n += 1
    return n


def acha_historico(doc):
    """Tabela do "Histórico de Versões", pelo cabeçalho completo.

    Exige as duas primeiras colunas ("Data" e "Versão") porque a tabela do Aceite
    também tem uma coluna "Data" — casar só por ela sobrescreve o quadro de assinatura.
    """
    for t in doc.tables:
        if not t.rows or len(t.rows[0].cells) < len(HISTORICO):
            continue
        cabecalho = [c.text.strip() for c in t.rows[0].cells]
        if cabecalho[0].startswith('Data') and cabecalho[1].startswith('Versão'):
            return t
    return None


def preenche_historico(doc, historico):
    """Preenche a tabela "Histórico de Versões" (cabeçalho + uma linha por versão).

    As linhas em branco que sobram são mantidas — é onde as revisões seguintes do
    documento vão ser anotadas à mão.
    """
    tab = acha_historico(doc)
    if tab is None:
        return 0
    livres = tab.rows[1:]
    if len(historico) > len(livres):
        modelo = copy.deepcopy(livres[-1]._tr)
        for _ in range(len(historico) - len(livres)):
            tab._tbl.append(copy.deepcopy(modelo))
        livres = tab.rows[1:]
    for row, versao in zip(livres, historico):
        for tc, (chave, rotulo) in zip(row._tr.tc_lst, HISTORICO):
            if chave not in versao:
                raise ErroConteudo('historico: falta a %s de uma das versões' % rotulo)
            texto_celula(tc, str(versao[chave]), doc)
    return len(historico)


def marca_checkbox(cell, alvo):
    """Troca ☐ por ☒ na opção escolhida (alvo = texto que segue o checkbox)."""
    if not alvo:
        return
    for p in cell.paragraphs:
        for r in p.runs:
            if VAZIO in r.text and alvo in r.text:
                partes = r.text.split(VAZIO)
                novo = partes[0]
                for trecho in partes[1:]:
                    marca = MARCADO if trecho.lstrip().startswith(alvo) else VAZIO
                    novo += marca + trecho
                r.text = novo


def respiro(p_el, onde='fim'):
    """Insere uma quebra de linha (<w:br>, o Shift+Enter do Word) para separar blocos.

    Clona o primeiro run do parágrafo para herdar a formatação e troca o texto pela
    quebra. É assim que a separação entre blocos aparece nas MIT044 revisadas à mão.
    """
    if p_el is None or p_el.tag != W_P:
        return p_el
    runs = p_el.findall(W_R)
    if not runs:
        return p_el
    novo = copy.deepcopy(runs[0])
    for t in novo.findall(W_T):
        novo.remove(t)
    for br in novo.findall(qn('w:br')):
        novo.remove(br)
    etree.SubElement(novo, qn('w:br'))
    if onde == 'fim':
        runs[-1].addnext(novo)
    else:
        runs[0].addprevious(novo)
    return p_el


def quebra_pagina_antes(tabela):
    """Faz a tabela começar em página nova (o "Quebrar página antes" do Word).

    Vai no primeiro parágrafo da primeira célula: é assim que o Word marca a quebra
    para uma tabela. O template traz `pageBreakBefore w:val="0"` em vários parágrafos,
    então o valor existente é sobrescrito em vez de duplicado.
    """
    if tabela is None or not tabela.rows:
        return False
    p = tabela.rows[0].cells[0].paragraphs[0]._p
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(p, qn('w:pPr'))
        p.insert(0, pPr)
    quebra = pPr.find(qn('w:pageBreakBefore'))
    if quebra is None:
        quebra = etree.Element(qn('w:pageBreakBefore'))
        # no OOXML pageBreakBefore vem logo depois de pStyle/keepNext/keepLines
        anterior = None
        for tag in ('w:pStyle', 'w:keepNext', 'w:keepLines'):
            achado = pPr.find(qn(tag))
            if achado is not None:
                anterior = achado
        if anterior is not None:
            anterior.addnext(quebra)
        else:
            pPr.insert(0, quebra)
    quebra.set(qn('w:val'), '1')
    return True


def respiro_entre_secoes(blocos):
    """Uma linha de respiro antes de cada título de seção.

    A quebra vai no fim do último elemento da seção anterior — o mesmo lugar em que
    ela é feita à mão no Word. Quando a seção termina em tabela não há onde inserir
    o <w:br>, e o bloco é pulado.
    """
    n = 0
    for els in blocos:
        if els and els[-1].tag == W_P:
            respiro(els[-1], 'fim')
            n += 1
    return n


def aplica_respiro(elementos):
    """Separa visualmente os blocos: a quebra vai no fim do elemento anterior; quando
    o anterior é uma tabela (onde não cabe <w:br>), vai no início do próprio label."""
    for i, el in enumerate(elementos):
        if i == 0 or el.tag != W_P:
            continue
        if ''.join(el.itertext()).strip() not in TODOS_LABELS:
            continue
        anterior = elementos[i - 1]
        if anterior.tag == W_TBL:
            respiro(el, 'inicio')
        else:
            respiro(anterior, 'fim')
    return elementos


def insere_apos(ancora, elementos):
    cur = ancora
    for el in elementos:
        cur.addnext(el)
        cur = el
    return cur


# ------------------------------------------------------------------ conteudo
def itens_para_elementos(itens, protos, contexto):
    """Converte a lista do JSON em elementos <w:p>. Item pode ser string (parágrafo)
    ou objeto {"tipo": "p"|"bullet"|"num"|"label", "texto": "..."}."""
    out = []
    n_num = 0
    for item in itens:
        if isinstance(item, str):
            tipo, texto = 'p', item
        elif isinstance(item, dict):
            tipo, texto = item.get('tipo', 'p'), item.get('texto', '')
        else:
            raise ErroConteudo('item inválido em %s: %r' % (contexto, item))
        if not texto:
            raise ErroConteudo('item sem texto em ' + contexto)
        if tipo == 'bullet':
            out.append(clona_paragrafo(protos['bullet'], BULLET + texto))
        elif tipo == 'num':
            n_num += 1
            out.append(clona_paragrafo(protos['bullet'], '%d.  %s' % (n_num, texto)))
        elif tipo == 'label':
            out.append(clona_paragrafo(protos['label'], texto))
        elif tipo == 'p':
            out.append(clona_paragrafo(protos['body'], texto))
        else:
            raise ErroConteudo('tipo desconhecido "%s" em %s' % (tipo, contexto))
    return out


def bloco(protos, label, itens, formato, contexto):
    """Label em negrito + itens no formato indicado (bullet/num/p)."""
    els = [clona_paragrafo(protos['label'], label)]
    if isinstance(itens, str):
        itens = [itens]
    els += itens_para_elementos(
        [{'tipo': formato, 'texto': i} if isinstance(i, str) else i for i in itens],
        protos, contexto)
    return els


def monta_execucao(doc, protos, exe):
    els = []
    for chave, label, formato in BLOCOS_EXECUCAO:
        if chave not in exe:
            raise ErroConteudo('execucao."%s" ausente no JSON' % chave)
        els += bloco(protos, label, exe[chave], formato, 'execucao.' + chave)
    return els


def monta_customizacoes(doc, protos, cus):
    for obrig in ('periodicidade', 'onde_executada', 'funcionalidades',
                  'premissas_tecnicas', 'prototipo_tela', 'anexos'):
        if obrig not in cus:
            raise ErroConteudo('customizacoes."%s" ausente no JSON' % obrig)
    els = []

    # Periodicidade de execucao + tabela 4x2
    per = cus['periodicidade']
    marcada = per.get('marcada', 'sob_demanda')
    linhas = [['Forma de execução', 'Marcação']]
    for chave, padrao in (('sob_demanda', 'Execução sob demanda'),
                          ('job', 'Job / processamento agendado'),
                          ('continua', 'Execução contínua')):
        linhas.append([per.get(chave, padrao), 'X' if marcada == chave else ''])
    els.append(clona_paragrafo(protos['label'], 'Periodicidade de execução:'))
    els.append(clona_tabela(protos['periodicidade'], linhas, doc))

    # Onde sera executada + tabela 2x2
    onde = cus['onde_executada']
    els.append(clona_paragrafo(protos['label'], 'Onde será executada:'))
    els.append(clona_paragrafo(protos['body'], onde.get('texto', '')))
    els.append(clona_tabela(protos['rotina'], [
        ['Rotina (código provisório)', 'Descrição de menu'],
        [onde.get('rotina', 'A definir'), onde.get('menu', '')],
    ], doc))

    els += bloco(protos, 'Funcionalidades:', cus['funcionalidades'], 'bullet',
                 'customizacoes.funcionalidades')
    els += bloco(protos, 'Premissas e restrições técnicas:', cus['premissas_tecnicas'],
                 'bullet', 'customizacoes.premissas_tecnicas')
    els += bloco(protos, 'Protótipo de tela:', cus['prototipo_tela'], 'p',
                 'customizacoes.prototipo_tela')

    # Anexos + tabela
    linhas = [['Descrição', 'Observação']]
    for linha in cus['anexos']:
        if not isinstance(linha, (list, tuple)) or len(linha) != 2:
            raise ErroConteudo('cada anexo deve ser um par [descrição, observação]')
        linhas.append(list(linha))
    els.append(clona_paragrafo(protos['label'], 'Anexos:'))
    els.append(clona_tabela(protos['anexos'], linhas, doc))
    return els


# ------------------------------------------------------------------ principal
def valida(conteudo):
    for chave in ('arquivo', 'capa', 'processo_atual', 'processo_proposto',
                  'parametrizacoes', 'execucao', 'customizacoes'):
        if chave not in conteudo:
            raise ErroConteudo('seção obrigatória ausente no JSON: ' + chave)
    arq = conteudo['arquivo']
    for chave in ('tag_cliente', 'titulo'):
        if not arq.get(chave):
            raise ErroConteudo('arquivo."%s" é obrigatório' % chave)
    for chave in ('nome_cliente', 'data'):
        if not conteudo['capa'].get(chave):
            raise ErroConteudo('capa."%s" é obrigatório' % chave)
    hist = conteudo.get('historico')
    if hist is not None and not isinstance(hist, list):
        raise ErroConteudo('"historico" deve ser uma lista de versões')


def historico_padrao(capa):
    """Primeira linha do histórico quando o JSON não traz "historico"."""
    return [{
        'data': capa.get('data', ''),
        'versao': '1.00',
        'autor': capa.get('responsavel_totvs') or capa.get('gerente_totvs', ''),
        'descricao': 'Emissão inicial do documento.',
    }]


def nome_saida(arq):
    titulo = arq['titulo'].strip()
    return '[%s] - Especificação da Customização - MIT044 - %s.docx' % (
        arq['tag_cliente'].strip(), titulo)


def main():
    ap = argparse.ArgumentParser(description='Gera MIT044 (.docx) a partir de um JSON.')
    ap.add_argument('json', help='arquivo JSON com o conteúdo')
    ap.add_argument('--saida', default='.', help='pasta de destino (padrão: atual)')
    ap.add_argument('--template', default=TEMPLATE_PADRAO, help='template .docx alternativo')
    ap.add_argument('--force', action='store_true',
                    help='sobrescreve o destino (gera .bak antes)')
    ap.add_argument('--sem-respiro', action='store_true',
                    help='não insere as quebras de linha que separam os blocos')
    ap.add_argument('--numeracao-uniforme', action='store_true',
                    help='põe os 5 títulos de seção na mesma lista numerada '
                         '(corrige a numeração quebrada herdada dos documentos originais)')
    args = ap.parse_args()

    with open(args.json, encoding='utf-8') as fh:
        conteudo = json.load(fh)
    valida(conteudo)

    if not os.path.isfile(args.template):
        raise SystemExit('template não encontrado: ' + args.template)

    destino = os.path.join(os.path.abspath(args.saida), nome_saida(conteudo['arquivo']))
    if os.path.exists(destino):
        if not args.force:
            raise SystemExit('arquivo já existe (use --force para sobrescrever):\n  ' + destino)
        shutil.copy2(destino, destino + '.bak')
        print('backup:', destino + '.bak')

    protos = carrega_prototipos()
    doc = docx.Document(args.template)
    heads = acha_headings(doc)
    removidos = limpa_miolo(doc, heads)
    if removidos:
        print('miolo do template limpo: %d elementos' % removidos)
    quebras = limpa_quebras(heads)
    if quebras:
        print('quebras de linha herdadas nos títulos removidas: %d' % quebras)

    if args.numeracao_uniforme:
        print('numeração dos títulos uniformizada: %d seções'
              % uniformiza_numeracao(doc, heads))

    if corrige_sumario(doc):
        print('campo do sumário ajustado para seleção por nível de título')

    preenche_capa(doc, conteudo['capa'])
    versoes = conteudo.get('historico') or historico_padrao(conteudo['capa'])
    if preenche_historico(doc, versoes):
        print('histórico de versões: %d linha(s)' % len(versoes))
    instrucoes = remove_instrucoes(doc)
    if instrucoes:
        print('textos de orientação do template removidos: %d' % instrucoes)

    if quebra_pagina_antes(acha_tabela(doc, 'Dados da Customização', 1)):
        print('quadro "Dados da Customização" começa em página nova')

    els_atu = itens_para_elementos(conteudo['processo_atual'], protos, 'processo_atual')
    els_pro = itens_para_elementos(conteudo['processo_proposto'], protos, 'processo_proposto')
    els_par = itens_para_elementos(conteudo['parametrizacoes'], protos, 'parametrizacoes')
    els_exe = monta_execucao(doc, protos, conteudo['execucao'])
    els_cus = monta_customizacoes(doc, protos, conteudo['customizacoes'])
    if not args.sem_respiro:
        aplica_respiro(els_exe)
        aplica_respiro(els_cus)
        # respiro depois do título da seção e antes do Aceite (precedido de tabela)
        respiro(heads['Execução'], 'fim')
        respiro(heads['Customizações'], 'fim')
        respiro(heads['Aceite'], 'inicio')
        # e antes de cada título de seção, no fim da seção anterior
        respiro_entre_secoes([els_atu, els_pro, els_par, els_exe])
    insere_apos(heads['Processo Atual'], els_atu)
    insere_apos(heads['Processo Proposto'], els_pro)
    insere_apos(heads['Parametrizações'], els_par)
    insere_apos(heads['Execução'], els_exe)
    insere_apos(heads['Customizações'], els_cus)

    sobras = limpa_placeholders(doc)
    if sobras:
        print('placeholders {{...}} do template limpos: %d' % sobras)

    os.makedirs(os.path.dirname(destino), exist_ok=True)
    doc.save(destino)
    print('gerado:', destino)
    print('LEMBRETE: abra no Word e atualize o sumário (clique no sumário e tecle F9).')


if __name__ == '__main__':
    try:
        main()
    except ErroConteudo as e:
        raise SystemExit('ERRO de conteúdo: %s' % e)
