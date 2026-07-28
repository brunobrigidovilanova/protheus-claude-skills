# -*- coding: utf-8 -*-
"""Confere a estrutura de uma MIT044 gerada, comparando com um documento de referencia.

Uso:
    python valida_mit044.py <gerado.docx> [--ref <referencia.docx>]

Sai com codigo 1 se alguma verificacao falhar.
"""
import sys, os, re, argparse, zipfile
sys.stdout.reconfigure(encoding='utf-8')
import docx
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

# titulos que precisam existir, nesta ordem. "Ambientação" e "Histórico de Versões"
# nao entram: no template oficial o primeiro e' formatado direto (sem estilo de titulo)
# e o segundo so' existe nas versoes mais novas do documento.
ESQUELETO = ['Sumário', 'Especificação da Customização', 'Processo Atual',
             'Processo Proposto', 'Parametrizações', 'Execução', 'Customizações', 'Aceite']
SECOES = ['Processo Atual', 'Processo Proposto', 'Parametrizações', 'Execução', 'Customizações']
LABELS = [
    'Objetivos do negócio:', 'Fluxo do processo:', 'Premissas e Restrições:',
    'Plano de teste e cenários esperados:', 'Rastreabilidade / dependência com outra MIT044:',
    'Periodicidade de execução:', 'Onde será executada:', 'Funcionalidades:',
    'Premissas e restrições técnicas:', 'Protótipo de tela:', 'Anexos:',
]
# tabelas obrigatorias, identificadas pela 1a celula preenchida, na ordem do documento
TABELAS = ['Nome do cliente', 'Dados da Customização', 'Forma de execução',
           'Rotina (código provisório)', 'Descrição', 'Aprovado por']
PLACEHOLDER = re.compile(r'\{\{[^{}]*\}\}')
INSTRUCAO = re.compile(r'^<.*>\.?$', re.S)

falhas = []
total = [0]


def ok(cond, msg, detalhe=''):
    total[0] += 1
    print(('  OK   ' if cond else '  FALHA') + '  ' + msg + (('  -> ' + detalhe) if detalhe and not cond else ''))
    if not cond:
        falhas.append(msg)


def rotulo_tabela(t):
    """Texto que identifica a tabela: 1a celula nao vazia da 1a linha (ou da 2a,
    quando a 1a linha e' mesclada e vazia, como na capa)."""
    for row in t.rows[:2]:
        for cell in row.cells:
            txt = cell.text.strip()
            if txt:
                return txt
    return ''


def acha_tabela(doc, prefixo):
    for t in doc.tables:
        if rotulo_tabela(t).startswith(prefixo):
            return t
    return None


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('docx')
    ap.add_argument('--ref', default=None, help='MIT044 de referência para comparação')
    args = ap.parse_args()

    doc = docx.Document(args.docx)
    print('Documento:', os.path.basename(args.docx))
    print('Tamanho  :', os.path.getsize(args.docx), 'bytes\n')

    # 1. esqueleto de titulos
    titulos = [(p.style.name, p.text.strip()) for p in doc.paragraphs
               if p.text.strip() and p.style.name in ('Title', 'Heading 1', 'Heading 2')]
    textos = [t for _, t in titulos]
    na_ordem = [t for t in textos if t in ESQUELETO]
    ok(na_ordem == ESQUELETO, 'títulos obrigatórios na ordem canônica', repr(textos))
    estilos_secao = {t: e for e, t in titulos}
    erradas = [s for s in SECOES if estilos_secao.get(s) != 'Heading 2']
    ok(not erradas, 'as cinco seções são Heading 2', 'fora do padrão: ' + repr(erradas))

    # 2. tabelas
    rotulos = [rotulo_tabela(t) for t in doc.tables]
    presentes = [r for r in rotulos if any(r.startswith(p) for p in TABELAS)]
    ok(len(presentes) == len(TABELAS)
       and all(r.startswith(p) for r, p in zip(presentes, TABELAS)),
       'tabelas obrigatórias presentes e na ordem esperada', repr(rotulos))

    capa = acha_tabela(doc, 'Nome do cliente')
    if capa is not None:
        ok(all(row.cells[0].text.split(':', 1)[-1].strip()
               for row in capa.rows if row.cells[0].text.strip()),
           'capa preenchida (sem rótulo órfão)')
    periodicidade = acha_tabela(doc, 'Forma de execução')
    if periodicidade is not None:
        marc = [r for r in periodicidade.rows if r.cells[1].text.strip() == 'X']
        ok(len(marc) == 1, 'exatamente uma forma de execução marcada com X')

    dados = acha_tabela(doc, 'Dados da Customização')
    if dados is not None:
        pPr = dados.rows[0].cells[0].paragraphs[0]._p.find(qn('w:pPr'))
        quebra = pPr.find(qn('w:pageBreakBefore')) if pPr is not None else None
        ok(quebra is not None and quebra.get(qn('w:val')) not in ('0', 'false'),
           'quadro "Dados da Customização" começa em página nova')

    aceite = acha_tabela(doc, 'Aprovado por')
    if aceite is not None:
        ok(len(aceite.rows) == 2, 'quadro de aceite com uma linha de assinatura',
           '%d linhas' % len(aceite.rows))

    # historico de versoes (so' nos templates que o trazem). O cabecalho e' conferido
    # nas duas primeiras colunas: a tabela do Aceite tambem tem uma coluna "Data".
    historico = None
    for t in doc.tables:
        cabecalho = [c.text.strip() for c in t.rows[0].cells] if t.rows else []
        if len(cabecalho) >= 4 and cabecalho[0].startswith('Data') \
                and cabecalho[1].startswith('Versão'):
            historico = t
            break
    if historico is not None:
        preenchidas = [r for r in historico.rows[1:]
                       if any(c.text.strip() for c in r.cells)]
        ok(len(preenchidas) >= 1, 'histórico de versões com ao menos uma linha',
           '%d linhas preenchidas' % len(preenchidas))

    # 3. bullets literais e labels
    paragrafos = [p.text for p in doc.paragraphs]
    n_bullets = sum(1 for t in paragrafos if t.startswith('•  '))
    ok(n_bullets > 0, 'bullets literais "•  " presentes', 'nenhum encontrado')
    faltando = [l for l in LABELS if l not in [t.strip() for t in paragrafos]]
    ok(not faltando, 'todos os labels de bloco presentes', 'faltando: ' + ', '.join(faltando))

    # labels em negrito
    negrito_ok = True
    for p in doc.paragraphs:
        if p.text.strip() in LABELS and p.runs:
            if not p.runs[0].bold:
                negrito_ok = False
    ok(negrito_ok, 'labels de bloco em negrito')

    # fluxo numerado (numeração literal "1.  ", não lista do Word)
    numerados = [t for t in paragrafos if re.match(r'^\d+\.\s\s', t.strip())]
    ok(len(numerados) >= 3, 'fluxo do processo numerado manualmente',
       '%d itens' % len(numerados))

    # 4. sobras do template: placeholders {{...}} e textos de orientação <...>
    tudo = paragrafos + [c.text for t in doc.tables for r in t.rows for c in r.cells]
    sobras = [t for t in tudo if PLACEHOLDER.search(t)]
    ok(not sobras, 'sem placeholders {{...}} do template', repr(sobras[:3]))
    guias = [t.strip() for t in paragrafos if t.strip() and INSTRUCAO.match(t.strip())]
    ok(not guias, 'sem textos de orientação do template', repr(guias[:2]))

    # 5. campo TOC
    toc = [e.text for e in doc.element.body.iter()
           if e.tag == qn('w:instrText') and e.text and 'TOC' in e.text]
    ok(bool(toc), 'campo de sumário (TOC) preservado')
    # com \t "Heading 1,1,..." o Word em português não acha nenhuma entrada
    ok(all('\\t' not in t for t in toc),
       'sumário selecionado por nível de título (\\o), não por nome de estilo', repr(toc))

    # 6. header/footer e imagens
    ok(len(doc.sections) == 1, 'documento com uma única seção')
    z = zipfile.ZipFile(args.docx)
    midias = [n for n in z.namelist() if n.startswith('word/media/')]
    fontes = [n for n in z.namelist() if n.endswith(('.odttf', '.ttf'))]
    ok(len(midias) >= 3, 'imagens do cabeçalho/rodapé preservadas',
       '%d encontradas' % len(midias))
    ok(len(fontes) > 0, 'fontes embutidas preservadas', 'nenhuma')
    partes_hf = [n for n in z.namelist() if re.match(r'word/(header|footer)\d*\.xml', n)]
    ok(len(partes_hf) >= 2, 'partes de cabeçalho e rodapé presentes',
       'encontradas: %r' % partes_hf)

    # 7. comparacao com referencia
    if args.ref:
        ref = docx.Document(args.ref)
        t_ref = [p.text.strip() for p in ref.paragraphs
                 if p.text.strip() and p.style.name in ('Title', 'Heading 1', 'Heading 2')]
        ok(textos == t_ref, 'esqueleto idêntico ao da MIT044 de referência',
           '%r vs %r' % (textos, t_ref))
        ok(len(doc.tables) == len(ref.tables), 'mesmo número de tabelas da referência',
           '%d vs %d' % (len(doc.tables), len(ref.tables)))
        est_ref = {t.style.name if t.style else None for t in ref.tables}
        est_ger = {t.style.name if t.style else None for t in doc.tables}
        ok(est_ger == est_ref, 'mesmos estilos de tabela da referência',
           '%r vs %r' % (est_ger, est_ref))

    if falhas:
        print('\n%d de %d verificações falharam: %s'
              % (len(falhas), total[0], '; '.join(falhas)))
    else:
        print('\nTUDO OK — %d verificações' % total[0])
    return 1 if falhas else 0


if __name__ == '__main__':
    sys.exit(main())
